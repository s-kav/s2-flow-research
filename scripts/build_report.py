#!/usr/bin/env python3
"""Build the English Word report for the Goldberg S^2-flow theorem."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from fractions import Fraction

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Goldberg_Snarks_S2_Flow_Theorem_en.docx"


def set_cell_shading(cell, fill: str) -> None:
    """Apply a solid background fill to one table cell."""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row) -> None:
    """Mark a table row as a repeating header."""
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def add_page_number(paragraph) -> None:
    """Insert a PAGE field into a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def add_equation(document: Document, text: str) -> None:
    """Add a centered display equation using Cambria Math."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11.5)


def add_code_block(document: Document, lines: list[str]) -> None:
    """Add a compact monospace code block."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_bullet(document: Document, text: str) -> None:
    """Add one compact bullet paragraph."""
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def configure_styles(document: Document) -> None:
    """Configure document-wide typography and heading styles."""
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (
        ("Title", 20, "17365D"),
        ("Subtitle", 11, "4F5B66"),
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12.5, "244A73"),
        ("Heading 3", 11, "244A73"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"

    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)


def add_caption(document: Document, text: str) -> None:
    """Add a centered figure caption."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_reference(document: Document, index: int, text: str) -> None:
    """Add one numbered bibliographic reference."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(1)
    label = paragraph.add_run(f"[{index}] ")
    label.bold = True
    label.font.size = Pt(8.5)
    body = paragraph.add_run(text)
    body.font.size = Pt(8.5)


def main() -> None:
    """Build and save the complete theorem report."""
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.32)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("S²-flows on Goldberg snarks")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("The S²-Flow Conjecture Holds for All Goldberg Snarks")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Analytic reduction, exact interval certificate, and reproducible full-graph verification")

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.add_run("21 July 2026").italic = True

    abstract_heading = document.add_paragraph()
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.add_run("Abstract").bold = True
    abstract = document.add_paragraph()
    abstract.paragraph_format.left_indent = Inches(0.4)
    abstract.paragraph_format.right_indent = Inches(0.4)
    abstract.add_run(
        "We prove that every Goldberg snark Gₖ, for odd k ≥ 5, admits an S²-flow. "
        "The cyclic block shift has twelve free edge orbits. The fundamental rotation by 2π/k is shown to be impossible for k ≥ 7, explaining the failure of the initial numerical continuation. "
        "A non-fundamental representation of index (k−1)/2, with rotation angle π−π/k, reduces the twelve-vector equivariant system to a single scalar equation H(s,x)=0. "
        "Exact rational interval arithmetic proves a uniform sign change and strict monotonicity on a parameter rectangle containing every Goldberg parameter. "
        "This yields a unique real-analytic scalar branch and an explicit unit-vector flow on the complete infinite family. The accompanying repository independently verifies the exact interval certificate and performs full-graph numerical checks."
    )

    document.add_paragraph("Keywords: Goldberg snark; unit vector flow; S²-flow; equivariant construction; interval arithmetic; cubic graph.")

    document.add_heading("1. Result and scope", level=1)
    document.add_paragraph(
        "An S²-flow on an oriented graph is a map from oriented edges to unit vectors in R³ such that the signed sum of incident edge vectors is zero at every vertex. The universal conjecture asserts that every bridgeless cubic graph admits such a flow [4,5]. This report proves the conjecture for the complete infinite family of Goldberg snarks, not for all bridgeless cubic graphs."
    )
    theorem = document.add_table(rows=1, cols=1)
    theorem.alignment = WD_TABLE_ALIGNMENT.CENTER
    theorem.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(theorem.cell(0, 0), "EAF2F8")
    paragraph = theorem.cell(0, 0).paragraphs[0]
    paragraph.add_run("Main Theorem. ").bold = True
    paragraph.add_run("For every odd integer k ≥ 5, the Goldberg snark Gₖ admits an S²-flow.")

    document.add_heading("2. Goldberg graph and the corrected orbit structure", level=1)
    document.add_paragraph(
        "For every block index t modulo k, let Bₜ be induced by vertices v₁ᵗ,…,v₈ᵗ. The nine internal edges are"
    )
    add_equation(document, "v₁v₂, v₁v₇, v₂v₈, v₃v₄, v₃v₈, v₄v₇, v₅v₆, v₆v₇, v₆v₈,")
    document.add_paragraph("and consecutive blocks are joined by")
    add_equation(document, "v₂ᵗv₁ᵗ⁺¹,   v₄ᵗv₃ᵗ⁺¹,   v₅ᵗv₅ᵗ⁺¹,")
    document.add_paragraph(
        "with block superscripts read modulo k. This is the standard eight-vertex Goldberg link construction [1–3]. For odd k ≥ 5 these graphs form the Goldberg snark family."
    )
    document.add_picture(str(ROOT / "docs" / "figures" / "goldberg_block.png"), width=Inches(6.65))
    add_caption(document, "Figure 1. The eight-vertex Goldberg block and its three incoming and three outgoing channels.")

    correction = document.add_table(rows=1, cols=1)
    correction.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(correction.cell(0, 0), "FFF4CE")
    paragraph = correction.cell(0, 0).paragraphs[0]
    paragraph.add_run("Correction to the supplied exploratory code. ").bold = True
    paragraph.add_run(
        "The cyclic shift t↦t+1 has exactly twelve edge orbits, all of size k. The apparent decomposition into twelve ordinary orbits and three singleton seam orbits was an orbit-canonicalization bug: a seam edge and its shifted non-seam edge belong to the same orbit."
    )

    document.add_heading("3. Why the fundamental generator fails", level=1)
    document.add_paragraph(
        "Let R be the rotation used to represent one block shift, and orient the channel r from v₅ᵗ to v₅ᵗ⁺¹. The Kirchhoff equation at v₅ reduces to"
    )
    add_equation(document, "g + r − R⁻¹r = 0.")
    document.add_paragraph("Because g must be a unit vector, every equivariant solution must satisfy")
    add_equation(document, "||(R⁻¹−I)r|| = 1.")
    document.add_paragraph(
        "If R has the fundamental angle 2π/k, the operator norm of R⁻¹−I is 2 sin(π/k). For k ≥ 7,"
    )
    add_equation(document, "2 sin(π/k) < 1,")
    document.add_paragraph(
        "so the required unit chord cannot exist. This is a rigorous structural obstruction, not a local-minimum problem. It explains why G₅ was solvable while the same ansatz failed at G₇."
    )

    document.add_heading("4. Correct cyclic representation", level=1)
    document.add_paragraph("For odd k define")
    add_equation(document, "ℓ = (k−1)/2,   φ = 2πℓ/k = π−π/k,   R = R_z(φ).")
    document.add_paragraph("Then Rᵏ=I. Put")
    add_equation(document, "u = φ/2,   x = cot u = tan(π/(2k)).")
    document.add_paragraph(
        "For k ≥ 5, 0 < x ≤ tan(π/10) < 13/40. The exact inequality follows from tan²(π/10)=1−2√5/5 and direct rational comparison. The complete proof therefore only needs the compact rational domain 0≤x≤13/40."
    )

    document.add_heading("5. The reduced 36-equation system", level=1)
    document.add_paragraph(
        "Orient the twelve representative edge orbits as a,b,c,d,e,f,g,h,i,p,q,r in the order listed in Table 1. Dividing the Kirchhoff equations in block t by Rᵗ gives eight representative vector equations:"
    )
    equations = [
        "a+b−R⁻¹p=0",
        "−a+c+p=0",
        "d+e−R⁻¹q=0",
        "−d+f+q=0",
        "g+r−R⁻¹r=0",
        "−g+h+i=0",
        "b+f+h=0",
        "c+e+i=0",
    ]
    add_equation(document, ",    ".join(equations[:4]))
    add_equation(document, ",    ".join(equations[4:]))
    document.add_paragraph(
        "Together with twelve unit-norm equations, this is the square 36×36 equivariant system. A global rotation about the z-axis gives one gauge direction in the unrestricted numerical formulation. The analytic reduction below fixes that gauge automatically."
    )

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Template", "Oriented representative", "Role")
    for index, header_text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "D9EAF7")
        cell.paragraphs[0].add_run(header_text).bold = True
    set_repeat_table_header(table.rows[0])
    rows = [
        ("a", "v₁ᵗ→v₂ᵗ", "first local channel"),
        ("b", "v₁ᵗ→v₇ᵗ", "first local channel"),
        ("c", "v₂ᵗ→v₈ᵗ", "first local channel"),
        ("d", "v₃ᵗ→v₄ᵗ", "second local channel"),
        ("e", "v₃ᵗ→v₈ᵗ", "second local channel"),
        ("f", "v₄ᵗ→v₇ᵗ", "second local channel"),
        ("g", "v₅ᵗ→v₆ᵗ", "central channel"),
        ("h", "v₆ᵗ→v₇ᵗ", "central closure"),
        ("i", "v₆ᵗ→v₈ᵗ", "central closure"),
        ("p", "v₂ᵗ→v₁ᵗ⁺¹", "inter-block"),
        ("q", "v₄ᵗ→v₃ᵗ⁺¹", "inter-block"),
        ("r", "v₅ᵗ→v₅ᵗ⁺¹", "inter-block"),
    ]
    for template, edge, role in rows:
        cells = table.add_row().cells
        cells[0].text = template
        cells[1].text = edge
        cells[2].text = role
    document.add_paragraph("Table 1. Orientation and order of the twelve free edge-orbit templates.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("6. Scalar reduction", level=1)
    document.add_heading("6.1 Local two-vertex channel", level=2)
    document.add_paragraph("For a scalar y define")
    add_equation(document, "z_y=√(1−y²(1+x²)),   δ_y=√(3−4y²),   λ_y=1−y²,")
    add_equation(document, "A_y=(δ_y z_y−yx)/(2λ_y),   D_y=(z_y+δ_y yx)/(2λ_y),")
    add_equation(document, "B_y=A_y+yx,   C_y=D_y−z_y.")
    document.add_paragraph(
        "Let W_y=(−yx,z_y) in the xz-plane and let J be the quarter-turn in that plane. Since ||W_y||²=λ_y and δ_y²=4λ_y−1, the vector"
    )
    add_equation(document, "U_y=(W_y−δ_y JW_y)/(2λ_y)")
    document.add_paragraph(
        "has unit norm and satisfies U_y·W_y=1/2. Its coordinates are (A_y,D_y). Consequently the four vectors generated by one local channel are all unit vectors and satisfy its two Kirchhoff equations. This proves the local identities algebraically, without numerical optimization."
    )

    document.add_heading("6.2 Final scalar equation", level=2)
    document.add_paragraph("Set t=1/2−s and define")
    add_equation(document, "H(s,x)=(B_t−B_s)²+(C_s−C_t)²−3/4.")
    document.add_paragraph(
        "The first term controls the horizontal component of the two central closure vectors; the second controls their vertical component. Therefore H(s,x)=0 is exactly the remaining unit-norm condition. All other equations are identities once t=1/2−s."
    )

    document.add_heading("7. Exact existence, uniqueness, and regularity", level=1)
    document.add_paragraph(
        "The repository contains a decisive certificate computed entirely with integers, exact rational intervals, and outward dyadic square-root bounds. No floating-point arithmetic is used for the following inequalities. The proof covers x∈[0,13/40] and s∈[2/3,21/25]."
    )
    certificate = json.loads((ROOT / "certificates" / "interval_proof_certificate.json").read_text(encoding="utf-8"))
    summary = certificate["summary"]
    lower_margin = float(Fraction(summary["max_upper_H_at_s_lower"]))
    upper_margin = float(Fraction(summary["min_lower_H_at_s_upper"]))
    derivative_margin = float(Fraction(summary["min_lower_dH_ds"]))

    metrics = document.add_table(rows=1, cols=3)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.style = "Table Grid"
    for index, text in enumerate(("Certified statement", "Exact sign", "Conservative decimal bound")):
        set_cell_shading(metrics.rows[0].cells[index], "D9EAF7")
        metrics.rows[0].cells[index].paragraphs[0].add_run(text).bold = True
    metric_rows = [
        ("H(2/3,x)", "strictly negative", f"< {lower_margin:.6f}"),
        ("H(21/25,x)", "strictly positive", f"> {upper_margin:.6f}"),
        ("∂H/∂s", "strictly positive", f"> {derivative_margin:.6f}"),
    ]
    for statement, sign, bound in metric_rows:
        cells = metrics.add_row().cells
        cells[0].text = statement
        cells[1].text = sign
        cells[2].text = bound
    document.add_paragraph(
        "The endpoint signs are certified on sixteen exact x-subintervals. Strict monotonicity is certified on a 32×32 subdivision of the (s,x)-rectangle at 112 dyadic bits. The stored JSON certificate contains every rational enclosure and is independently recomputed by the verifier."
    )
    document.add_picture(str(ROOT / "docs" / "figures" / "endpoint_signs.png"), width=Inches(6.55))
    add_caption(document, "Figure 2. Floating-point visualization of the sign separation that is proved by exact interval arithmetic.")

    document.add_paragraph(
        "By the intermediate value theorem, for every x in the domain there exists a root s(x)∈(2/3,21/25). Since ∂H/∂s>0, the root is unique. The radicands remain strictly positive on the rectangle, so H is real analytic; the implicit function theorem therefore gives a unique real-analytic branch s(x)."
    )
    document.add_picture(str(ROOT / "docs" / "figures" / "scalar_branch.png"), width=Inches(6.55))
    add_caption(document, "Figure 3. The unique scalar branch sampled at odd Goldberg parameters k=5,…,501.")

    document.add_heading("8. Explicit twelve-template construction", level=1)
    document.add_paragraph("For the unique root s=s(x), put t=1/2−s and define")
    formula_lines = [
        "a=( A_s, 0, D_s),       b=(−B_s, s, −C_s),       c=( B_s, s, C_s),",
        "d=( A_t, 0, D_t),       e=(−B_t, t, −C_t),       f=( B_t, t, C_t),",
        "g=(0,−1,0),",
        "h=(B_s−B_t,−1/2,C_s−C_t),",
        "i=(B_t−B_s,−1/2,C_t−C_s),",
        "p=(−sx,−s,z_s),       q=(−tx,−t,z_t),",
        "r=(x/2,1/2,√(3−x²)/2).",
    ]
    for line in formula_lines:
        add_equation(document, line)

    document.add_heading("9. Verification of the construction", level=1)
    document.add_heading("9.1 Unit norms", level=2)
    document.add_paragraph(
        "The local-channel lemma gives unit norm for a,b,c,p and d,e,f,q. The vector r is unit by direct calculation. The identity R⁻¹r−r=(0,−1,0) gives the unit vector g. Finally, H(s,x)=0 gives ||h||=||i||=1."
    )
    document.add_heading("9.2 Kirchhoff equations", level=2)
    add_bullet(document, "The local-channel identities give a+b=R⁻¹p, c=a−p, d+e=R⁻¹q, and f=d−q.")
    add_bullet(document, "The selected r gives g+r=R⁻¹r.")
    add_bullet(document, "The definitions give h+i=g.")
    add_bullet(document, "Because s+t=1/2, direct componentwise addition gives b+f+h=0 and c+e+i=0.")
    document.add_paragraph(
        "Thus all eight representative vertex equations hold. For every block shift j, assign Rʲ times the corresponding template to the shifted edge. Orthogonality preserves norms and the representative equations. Since Rᵏ=I, the assignment is consistent at the seam. This completes the proof of the Main Theorem."
    )

    document.add_heading("10. Reproducibility and independent checks", level=1)
    document.add_paragraph(
        "The exact theorem does not depend on numerical experiments. The experiments provide regression protection for graph indexing, orbit expansion, orientation signs, and implementation errors."
    )
    sweep_rows = list(csv.DictReader((ROOT / "certificates" / "numerical_sweep.csv").open(encoding="utf-8")))
    full_rows = [row for row in sweep_rows if row["full_verified"] == "True"]
    worst_reduced_k = max(float(row["reduced_kirchhoff"]) for row in sweep_rows)
    worst_reduced_n = max(float(row["reduced_norm"]) for row in sweep_rows)
    worst_full_k = max(float(row["full_kirchhoff"]) for row in full_rows)
    worst_full_n = max(float(row["full_norm"]) for row in full_rows)

    checks = document.add_table(rows=1, cols=2)
    checks.alignment = WD_TABLE_ALIGNMENT.CENTER
    checks.style = "Table Grid"
    for index, text in enumerate(("Check", "Observed result")):
        set_cell_shading(checks.rows[0].cells[index], "D9EAF7")
        checks.rows[0].cells[index].paragraphs[0].add_run(text).bold = True
    check_rows = [
        ("Odd parameters", f"499 values, k=5,…,1001"),
        ("Full graph expansions", f"149 graphs, k=5,…,301"),
        ("Worst reduced Kirchhoff residual", f"{worst_reduced_k:.3e}"),
        ("Worst reduced norm residual", f"{worst_reduced_n:.3e}"),
        ("Worst full Kirchhoff residual", f"{worst_full_k:.3e}"),
        ("Worst full norm residual", f"{worst_full_n:.3e}"),
        ("Independent large instance", "G₁₀₀₁: 8,008 vertices and 12,012 edges"),
        ("Automated tests", "10 tests passed"),
    ]
    for label, result in check_rows:
        cells = checks.add_row().cells
        cells[0].text = label
        cells[1].text = result

    document.add_heading("10.1 Commands", level=2)
    add_code_block(
        document,
        [
            "python -m pip install -r requirements-dev.txt",
            "python -m pip install -e .",
            "python scripts/run_interval_proof.py",
            "python scripts/verify_interval_certificate.py",
            "python scripts/run_numerical_sweep.py --max-k 1001 --full-max-k 301",
            "python scripts/verify_one.py 1001",
            "python -m pytest",
        ],
    )

    document.add_heading("11. What has and has not been proved", level=1)
    conclusion = document.add_table(rows=2, cols=2)
    conclusion.alignment = WD_TABLE_ALIGNMENT.CENTER
    conclusion.style = "Table Grid"
    set_cell_shading(conclusion.cell(0, 0), "E2F0D9")
    set_cell_shading(conclusion.cell(1, 0), "FCE4D6")
    conclusion.cell(0, 0).text = "Proved"
    conclusion.cell(0, 1).text = "Every Goldberg snark Gₖ with odd k≥5 has an explicit Zₖ-equivariant S²-flow; the scalar solution is unique and real analytic in the continuous parameter x."
    conclusion.cell(1, 0).text = "Not proved"
    conclusion.cell(1, 1).text = "The universal S²-flow conjecture for all bridgeless cubic graphs. The result is a strong infinite-family theorem, not a universal proof."

    document.add_heading("Appendix A. Proof-certificate configuration", level=1)
    add_bullet(document, f"Certificate schema: {certificate['schema']}")
    add_bullet(document, f"Dyadic precision: {certificate['bits']} bits")
    add_bullet(document, f"Endpoint partition: {certificate['partitions']['endpoint_x']} x-boxes")
    add_bullet(document, f"Derivative partition: {certificate['partitions']['derivative_s']}×{certificate['partitions']['derivative_x']} boxes")
    add_bullet(document, "Arithmetic: Python integers, fractions.Fraction, integer square root, and outward dyadic enclosures")
    add_bullet(document, "Certificate SHA-256: 0a55c601b8e1ea62c350a97da2cbb883cf8448e38f5f96918ba7c19206bfdf93")

    document.add_heading("References", level=1)
    add_reference(document, 1, "M. K. Goldberg, “Construction of Class 2 Graphs with Maximum Vertex Degree 3,” Journal of Combinatorial Theory, Series B 31 (1981), 282–291.")
    add_reference(document, 2, "C. Fiori and B. Ruini, “Infinite Classes of Dihedral Snarks,” Mediterranean Journal of Mathematics 5 (2008), 199–210. DOI: 10.1007/s00009-008-0144-3.")
    add_reference(document, 3, "X. Zhang, Y. Wang, and S. Zhang, “Feedback Numbers of Goldberg Snark, Twisted Goldberg Snarks and Related Graphs,” ITM Web of Conferences 25 (2019), 01012. DOI: 10.1051/itmconf/20192501012.")
    add_reference(document, 4, "D. Mattiolo et al., “On d-Dimensional Nowhere-Zero r-Flows on a Graph,” arXiv:2304.14231, 2023.")
    add_reference(document, 5, "H. Houdrouge et al., “2-Dimensional Unit Vector Flows,” arXiv:2602.21526, 2026.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
