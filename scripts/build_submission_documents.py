"""Build the journal manuscript and supplementary DOCX files."""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript"
FIG = OUT / "figures"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure(doc: Document, *, supplementary: bool = False) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62 if not supplementary else 0.7)
    sec.bottom_margin = Inches(0.62 if not supplementary else 0.7)
    sec.left_margin = Inches(0.68 if not supplementary else 0.75)
    sec.right_margin = Inches(0.68 if not supplementary else 0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.5 if not supplementary else 10)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.0
    for sty in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        styles[sty].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(17)
    styles["Heading 1"].font.size = Pt(12.5)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.space_before = Pt(7)
    styles["Heading 1"].paragraph_format.space_after = Pt(3)
    styles["Heading 2"].font.size = Pt(10.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].paragraph_format.space_before = Pt(5)
    styles["Heading 2"].paragraph_format.space_after = Pt(2)
    styles["Heading 3"].font.size = Pt(9.5)
    styles["Heading 3"].font.italic = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("S. Kavun | Spherical unit flows on snark families")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(17)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sergii Kavun")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Independent Researcher")
    r.italic = True
    r.font.size = Pt(9.5)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(9)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(text)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(9.5)


def add_theorem(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    r = p.add_run(label + ". ")
    r.bold = True
    r.italic = True
    p.add_run(text)


def add_figure(doc: Document, path: Path, caption: str, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(3)
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)


def add_references(doc: Document, refs: list[str]) -> None:
    doc.add_heading("References", level=1)
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(f"[{i}] ").bold = True
        p.add_run(ref)


def build_main() -> Path:
    doc = Document()
    configure(doc)
    add_title(doc, "Equivariant Unit-Sphere Flows on Flower and Goldberg Snarks")
    add_label_paragraph(doc, "Abstract.",
        "An S²-flow assigns a unit vector in R³ to each oriented edge and satisfies Kirchhoff conservation at every vertex. The conjecture that every bridgeless cubic graph admits such a flow remains open. We prove it for two classical infinite families of snarks. First, every Isaacs flower snark J_n, for odd n at least 5, admits a Z_n-equivariant S²-flow. A six-orbit reduction yields an 18-variable system, which is eliminated to one scalar equation whose endpoint signs and monotonicity give an analytic solution over the entire parameter interval. Second, every Goldberg snark G_k, for odd k at least 5, admits a Z_k-equivariant S²-flow. The fundamental cyclic representation is obstructed for k at least 7; the representation of index (k-1)/2 instead reduces the twelve-orbit system to a scalar equation. Exact rational interval arithmetic certifies a uniform sign change and strict derivative bound. Both constructions are explicit after solving their scalar equation. A reproducibility package independently checks the exact certificates and expands the flows on complete graphs.")
    add_label_paragraph(doc, "Keywords.", "snark; unit-vector flow; S²-flow; flower snark; Goldberg snark; equivariant construction; interval arithmetic")
    add_label_paragraph(doc, "Mathematics Subject Classification.", "05C21, 05C15, 05C30")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Nowhere-zero flows translate global graph structure into local conservation laws. Unit-vector flows replace scalar or group values by points on a sphere and thereby introduce geometric degrees of freedom without weakening Kirchhoff conservation. For an orientation of a graph G, an S²-flow is a map φ:E(G)→R³ such that ||φ(e)||=1 for every edge and the signed sum of incident values is zero at every vertex. The current universal conjecture asks whether every bridgeless cubic graph has an S²-flow [4–7].")
    doc.add_paragraph(
        "The conjecture is especially natural on snarks, where ordinary Tait colourings fail. Flower snarks, introduced by Isaacs [1], and Goldberg snarks, arising from Goldberg's class-2 construction [2], are canonical infinite test families. Their repeating block structures suggest equivariant flows, but symmetry is useful only after the correct cyclic representation has been identified. The two families exhibit complementary mechanisms: the flower construction uses six edge orbits and a Möbius-type indexing, while the Goldberg construction uses twelve free edge orbits and requires a non-fundamental representation of the block shift.")
    doc.add_paragraph("Our contributions are the following.")
    for item in [
        "A complete analytic proof that every flower snark J_n, with n odd and n≥5, admits an exact Z_n-equivariant S²-flow.",
        "A complete proof that every Goldberg snark G_k, with k odd and k≥5, admits an exact Z_k-equivariant S²-flow, including an exact rational interval certificate uniform in k.",
        "A structural explanation of two failed natural ansatzes: strict order-2n symmetry for flower snarks, and the fundamental rotation 2π/k for Goldberg snarks when k≥7.",
        "An independently checkable repository that distinguishes analytic proof, exact computer-assisted proof, finite exact certificates, and numerical validation."
    ]:
        p=doc.add_paragraph(style=None); p.style=doc.styles['Normal']; p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.first_line_indent=Inches(-0.12); p.add_run("• "); p.add_run(item)
    doc.add_paragraph(
        "The article is organised so that the two family proofs are self-contained at theorem level. Extended algebra, certificate formats, computational campaigns, and repository audit findings are placed in the Supplementary Material.")

    doc.add_heading("2. Preliminaries", level=1)
    doc.add_paragraph(
        "Fix an arbitrary orientation of a finite graph G. For v∈V(G), let δ⁺(v) and δ⁻(v) denote the outgoing and incoming edges. An S²-flow is an assignment φ:E(G)→S²⊂R³ satisfying")
    add_equation(doc, "Σ_{e∈δ⁺(v)} φ(e) − Σ_{e∈δ⁻(v)} φ(e) = 0  for every v∈V(G).")
    doc.add_paragraph(
        "Reversing an edge and negating its vector preserves the condition, so existence is orientation independent. At a cubic vertex, three unit vectors sum to zero exactly when they lie in a plane and form mutual angles 2π/3. This local geometric fact underlies the spherical-immersion viewpoint of recent work [7].")
    add_theorem(doc, "Lemma 2.1 (Equivariant expansion)",
        "Let a cyclic automorphism g of order m act freely on the oriented edge set, without edge inversion. If R∈SO(3) satisfies R^m=I and template vectors satisfy Kirchhoff's law at one vertex from each orbit, then φ(g^j e)=R^jφ(e) defines a flow on the whole graph. If the templates are unit vectors, the expanded flow is an S²-flow.")
    doc.add_paragraph(
        "The proof is immediate: incidence at g^jv is the image of incidence at v, so its Kirchhoff sum is R^j times the representative sum; orthogonality preserves norms and R^m=I makes the seam consistent.")

    doc.add_heading("3. Flower snarks", level=1)
    doc.add_heading("3.1. Six-orbit reduction", level=2)
    doc.add_paragraph(
        "Let J_n be the canonical flower snark for odd n≥5. Use hub vertices c_i, outer vertices t_i, and strand vertices w_j, with indices in Z_n or Z_{2n}. The two-block shift g has order n and acts freely. With R=R_z(2u), u=π/n, an equivariant flow is determined by six templates H,T,F,G,P,Q corresponding to the six edge orbits shown in Figure 1.")
    add_figure(doc, FIG/"flower_fundamental_cell.png", "Figure 1. Fundamental orbit representatives for the flower-snark reduction.", 5.6)
    doc.add_paragraph("Let A=R^{(n−1)/2}=R_z(π−u). Kirchhoff conservation at four representative vertex orbits gives")
    add_equation(doc, "H+F+AG=0,   T−H−AT=0,   P−F−R⁻¹Q=0,   Q−G−P=0,")
    doc.add_paragraph("together with ||H||=||T||=||F||=||G||=||P||=||Q||=1. The z-row of the second vector equation is the negative of the sum of the z-rows of the other three equations. Fixing the remaining rotation gauge by H_y=0 yields a square 18×18 system.")

    doc.add_heading("3.2. Scalar cascade", level=2)
    doc.add_paragraph(
        "The system can be eliminated without approximation. The z-row forces H_z=0; we choose H=(1,0,0). The equation (I−A)T=H fixes T in the xy-plane and its norm fixes T_z. Write G'=AG=(-1/2,η,ζ) with η²+ζ²=3/4; then F=(-1/2,−η,−ζ). Combining the last two Kirchhoff equations gives (I−R⁻¹)P=F+R⁻¹G and Q=G+P. Put m=η/sin u and choose ζ=−sqrt(3/4−η²). The planar part of P has squared norm")
    add_equation(doc, "ρ²(m)=1/[8(1+cos u)] + [m²(1+cos u)+m]/2.")
    doc.add_paragraph("Its admissible interval is [m_−,m_+], where")
    add_equation(doc, "m_±=[−1 ± 2sqrt(2(1+cos u))]/[2(1+cos u)].")
    doc.add_paragraph("The remaining norm equation ||Q||=1 is equivalent to h(m,u)=0, with")
    add_equation(doc, "h=3/8 − (m² sin²u)/2 − m/2 − sqrt(3/4−m²sin²u) sqrt(1−ρ²(m)).")

    add_theorem(doc, "Theorem 3.1 (Flower Snark Theorem)", "Every flower snark J_n with odd n≥5 admits a Z_n-equivariant S²-flow.")
    doc.add_paragraph("Proof. Fix u∈(0,π/5]. The factorisation")
    add_equation(doc, "ρ²−1=[(1+cos u)/2](m−m_−)(m−m_+)")
    doc.add_paragraph(
        "shows that h is real and continuous on [m_−,0]. On that interval |m sin u|<sqrt(3)/2, so the second square root is also real. At m=m_−, the factor sqrt(1−ρ²) vanishes and")
    add_equation(doc, "h(m_−,u)=3/8−m_−²sin²u/2−m_−/2>0.")
    doc.add_paragraph("At m=0,")
    add_equation(doc, "h(0,u)=3/8−(sqrt(3)/2)sqrt(1−1/[8(1+cos u)])<0,")
    doc.add_paragraph(
        "because u≤π/5 gives 1/[8(1+cos u)]<1/12. The intermediate value theorem gives m*(u)∈(m_−,0). Substitution into the cascade produces six unit templates satisfying the representative equations, and Lemma 2.1 expands them to J_n. ∎")
    doc.add_paragraph(
        "A derivative calculation shows h_m<0 on the same interval, hence the normalised root is unique. Since h is analytic and h_m does not vanish, the implicit function theorem yields a real-analytic branch in u. The detailed inequality chain is supplied in Supplementary Section S2.")

    doc.add_heading("4. Goldberg snarks", level=1)
    doc.add_heading("4.1. Graph and orbit structure", level=2)
    doc.add_paragraph(
        "For t∈Z_k, let B_t contain vertices v_1^t,…,v_8^t and internal edges v_1v_2, v_1v_7, v_2v_8, v_3v_4, v_3v_8, v_4v_7, v_5v_6, v_6v_7, v_6v_8. Consecutive blocks are joined by v_2^tv_1^{t+1}, v_4^tv_3^{t+1}, and v_5^tv_5^{t+1}, with indices modulo k. This is the eight-vertex Goldberg-link construction appearing in the literature [2,3]. The block shift has exactly twelve free edge orbits, each of size k.")
    add_figure(doc, FIG/"goldberg_block.png", "Figure 2. One Goldberg block and its three cyclic channels.", 5.2)

    doc.add_heading("4.2. Obstruction to the fundamental representation", level=2)
    doc.add_paragraph(
        "Suppose one block shift is represented by the fundamental rotation R_z(2π/k). At v_5, the representative equation has the form g+r−R⁻¹r=0. Since g and r are unit vectors, a necessary condition is ||(R⁻¹−I)r||=1. The operator norm of R⁻¹−I is 2sin(π/k), which is less than 1 for k≥7. Thus the fundamental representation cannot support this equivariant ansatz. The successful representation has index ℓ=(k−1)/2 and angle")
    add_equation(doc, "φ=2πℓ/k=π−π/k.")
    doc.add_paragraph("Set x=cot(φ/2)=tan(π/(2k)); then 0<x≤tan(π/10)<13/40.")

    doc.add_heading("4.3. Scalar reduction", level=2)
    doc.add_paragraph("For a scalar y define")
    add_equation(doc, "z_y=sqrt(1−y²(1+x²)),  δ_y=sqrt(3−4y²),  λ_y=1−y²,")
    add_equation(doc, "A_y=(δ_y z_y−yx)/(2λ_y),  D_y=(z_y+δ_y yx)/(2λ_y),")
    add_equation(doc, "B_y=A_y+yx,  C_y=D_y−z_y.")
    doc.add_paragraph(
        "These expressions solve one two-vertex channel algebraically: the associated four vectors are unit and satisfy both local Kirchhoff equations. Put t=1/2−s. All remaining representative equations and eleven of the twelve norm equations become identities, while the last norm condition is")
    add_equation(doc, "H(s,x)=(B_t−B_s)²+(C_s−C_t)²−3/4=0.")

    add_theorem(doc, "Theorem 4.1 (Goldberg Snark Theorem)", "Every Goldberg snark G_k with odd k≥5 admits a Z_k-equivariant S²-flow.")
    doc.add_paragraph(
        "Proof. Exact rational interval arithmetic on the rectangle 0≤x≤13/40 and 2/3≤s≤21/25 gives")
    add_equation(doc, "H(2/3,x)<−0.281883,   H(21/25,x)>0.058561,   ∂H/∂s>1.220892.")
    doc.add_paragraph(
        "The decimal values are conservative displays of exact rational bounds stored in the certificate. Hence, for each admissible x, H has exactly one root s(x) in the bracket. All radicands are strictly positive on the rectangle, so H is analytic and the implicit function theorem gives an analytic branch. For the unique root, put t=1/2−s and define the twelve templates")
    add_equation(doc, "a=(A_s,0,D_s), b=(−B_s,s,−C_s), c=(B_s,s,C_s),")
    add_equation(doc, "d=(A_t,0,D_t), e=(−B_t,t,−C_t), f=(B_t,t,C_t),")
    add_equation(doc, "g=(0,−1,0), h=(B_s−B_t,−1/2,C_s−C_t),")
    add_equation(doc, "i=(B_t−B_s,−1/2,C_t−C_s),")
    add_equation(doc, "p=(−sx,−s,z_s), q=(−tx,−t,z_t), r=(x/2,1/2,sqrt(3−x²)/2).")
    doc.add_paragraph(
        "The local-channel identities prove the norms and four representative Kirchhoff equations; the definition of r gives g+r=R⁻¹r; h+i=g; s+t=1/2 gives the two central closure equations; and H=0 gives ||h||=||i||=1. Lemma 2.1 expands the templates consistently because R^k=I. ∎")

    doc.add_heading("5. Exact certification and computational validation", level=1)
    doc.add_paragraph(
        "The flower theorem is analytic. A separate chain of 1,510 rational Newton-Kantorovich boxes covers a large finite parameter interval and supplies an independent computer-rigorous check. Exact dyadic certificates are also stored for J_5 through J_41. The Goldberg proof contains one computer-assisted component: rational interval evaluation of endpoint signs and H_s on a compact rectangle. The verifier recomputes every outward-rounded bound from integer numerators and denominators and checks the certificate hash.")
    table=doc.add_table(rows=1, cols=3)
    table.style='Table Grid'
    hdr=table.rows[0].cells
    for i,t in enumerate(['Run','Coverage','Status']): hdr[i].text=t; set_cell_shading(hdr[i],'D9EAF7')
    rows=[
        ('Flower analytic verifier','all odd n≥5','symbolic/rational support for analytic proof'),
        ('Flower dyadic certificates','J_5,…,J_41','exact finite certificates'),
        ('Goldberg interval certificate','all odd k≥5','exact computer-assisted proof component'),
        ('Goldberg sweep','499 odd k≤1001','numerical regression validation'),
        ('Full Goldberg expansions','149 graphs through k=301','orientation and indexing validation'),
    ]
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    doc.add_paragraph(
        "Finite sweeps are not used to infer an infinite result. They are included to detect graph-construction, orientation, seam, and floating-point implementation errors. Complete commands and generated paths appear in RUNBOOK.md.")

    doc.add_heading("6. Comparison and implications", level=1)
    doc.add_paragraph(
        "Both proofs convert an infinite graph family into a compact one-parameter problem. The decisive feature is not cyclic symmetry alone, but a compatible representation of the cyclic action in SO(3). For flower snarks the index-one Z_n action is compatible after passing to the two-block shift. For Goldberg snarks, index one is obstructed by a unit-chord inequality and the high-frequency index (k−1)/2 is required. In both cases, local Kirchhoff equations can be solved first and global closure is reduced to one scalar equation.")
    doc.add_paragraph(
        "This suggests a reusable programme for other periodic snarks: determine free edge orbits; enumerate representations of the cyclic automorphism; reject representations by local chord or fixed-space constraints; eliminate local channels; and certify the remaining low-dimensional closure equation by analytic or interval methods. The present results do not imply the universal conjecture, but they enlarge the exact class of non-3-edge-colourable cubic graphs known to admit S²-flows.")

    doc.add_heading("7. Reproducibility, limitations, and release status", level=1)
    doc.add_paragraph(
        "The source archive contains exact certificates, independent verifiers, test cases, manuscript generators, and a complete runbook. The reviewed release resolves an original licence merge conflict, removes machine-specific paths and generated caches, corrects the default certificate location, and adds continuous integration. A public archival release should add the confirmed GitHub URL to CITATION.cff and mint a permanent DOI.")
    doc.add_paragraph(
        "The mathematical limitation is explicit: no claim is made for arbitrary bridgeless cubic graphs. Independent expert review is still required before publication, especially for the elimination steps and the uniform inequalities. The code supports verification but does not replace peer review.")

    doc.add_heading("8. Conclusion", level=1)
    doc.add_paragraph(
        "Every flower snark and every Goldberg snark in their standard odd-parameter families admits an exact equivariant S²-flow. The proofs exhibit two complementary scalar reductions and show how representation choice controls the success of a symmetry ansatz. Exact rational certification and full-graph regression checks make the computer-assisted components independently reproducible.")

    refs=[
        "R. Isaacs, Infinite families of nontrivial trivalent graphs which are not Tait colorable, American Mathematical Monthly 82 (1975), 221–239. doi:10.2307/2319844.",
        "M. K. Goldberg, Construction of class 2 graphs with maximum vertex degree 3, Journal of Combinatorial Theory, Series B 31 (1981), 282–291. doi:10.1016/0095-8956(81)90030-7.",
        "H. Fleischner, B. Bagheri Gh., and B. Klocker, Perfect pseudo-matchings in cubic graphs, Graphs and Combinatorics 40 (2024), Article 118. doi:10.1007/s00373-024-02844-y.",
        "C. Thomassen, Group flow, complex flow, unit vector flow, and the (2+epsilon)-flow conjecture, Journal of Combinatorial Theory, Series B 108 (2014), 81–91. doi:10.1016/j.jctb.2014.02.012.",
        "Y. Wang, J. Cheng, R. Luo, and C.-Q. Zhang, Vector flows and integer flows, SIAM Journal on Discrete Mathematics 29 (2015), 2166–2178. doi:10.1137/151006329.",
        "D. Mattiolo, G. Mazzuoccolo, J. Rajník, and G. Tabarelli, On d-dimensional nowhere-zero r-flows on a graph, European Journal of Mathematics 9 (2023), Article 101. doi:10.1007/s40879-023-00694-1.",
        "D. Mattiolo, G. Mazzuoccolo, J. Rajník, and G. Tabarelli, Geometric description of d-dimensional flows of a graph, Australasian Journal of Combinatorics 94 (2026), 376–384.",
        "H. Houdrouge, B. Miraftab, and P. Morin, 2-dimensional unit vector flows, arXiv:2602.21526 (2026).",
    ]
    add_references(doc, refs)
    p=doc.add_paragraph()
    p.add_run("Data and code availability. ").bold=True
    p.add_run("The reviewed reproducibility package accompanying this submission contains all source code, exact certificates, independent verifiers, generated tables, and manuscript sources. Release metadata and the public repository URL should be frozen at acceptance or preprint deposit.")
    path=OUT/'S2_Flows_Flower_Goldberg_Snarks.docx'
    doc.save(path)
    return path


def build_supplement() -> Path:
    doc=Document(); configure(doc, supplementary=True)
    add_title(doc, "Supplementary Material: Equivariant Unit-Sphere Flows on Flower and Goldberg Snarks",
              "Extended derivations, exact certificates, computational validation, and repository audit")
    doc.add_heading("S1. Scope and epistemic classification", level=1)
    doc.add_paragraph(
        "This supplement records derivations and verification details omitted from the main article. Statements are classified as analytic proofs, exact computer-assisted proofs, exact finite certificates, numerical validation, or repository-engineering observations. Numerical residuals are never used as substitutes for an infinite theorem.")
    table=doc.add_table(rows=1,cols=2); table.style='Table Grid'
    for i,t in enumerate(['Category','Meaning']): table.rows[0].cells[i].text=t; set_cell_shading(table.rows[0].cells[i],'D9EAF7')
    for a,b in [
        ('Analytic proof','Finite symbolic argument valid for every admissible parameter.'),
        ('Exact computer-assisted proof','A finite rational/integer certificate rechecked by a small independent verifier.'),
        ('Exact finite certificate','Existence for a specified graph or compact parameter box.'),
        ('Numerical validation','Floating-point regression check; useful for error detection, not proof.'),
    ]:
        c=table.add_row().cells;c[0].text=a;c[1].text=b

    doc.add_heading("S2. Flower-snark derivation", level=1)
    doc.add_heading("S2.1. Canonical equations and dependency", level=2)
    doc.add_paragraph(
        "With H,T,F,G,P,Q and A=R_z(π−u), R=R_z(2u), the representative equations are E1–E4 as stated in the article. Rotations fix the z-axis, hence (E1+E3+E4)_z=H_z and (E2)_z=−H_z. The full system is therefore recovered from E1, E3, E4, the xy-rows of E2, six norms, and H_y=0.")
    doc.add_heading("S2.2. Closed-form cascade", level=2)
    doc.add_paragraph(
        "The steps H=(1,0,0), solution of (I−A)T=H, circle parameterisation of AG, planar inversion of I−R⁻¹, and Q=G+P are equivalences. No numerical solver is required. The exact identity")
    add_equation(doc, "ρ²=[(1+cos u)/2](m−m₀)²,  m₀=−1/[2(1+cos u)],")
    doc.add_paragraph("makes the admissible interval transparent and prevents the parameter interval from collapsing as u→0.")
    doc.add_heading("S2.3. Endpoint inequalities", level=2)
    doc.add_paragraph(
        "At m=m_− the radial square root is zero. The remaining expression is positive because m_−<0 and η²<3/4. At m=0, η=0 and the negative product term dominates 3/8 uniformly for u≤π/5. These exact inequalities establish existence by the intermediate value theorem.")
    doc.add_heading("S2.4. Monotonicity and analytic continuation", level=2)
    doc.add_paragraph(
        "Differentiating h gives h_m=−m sin²u−1/2+ζ_m w+ζ w_m, where ζ<0, ζ_m=m sin²u/|ζ|, w=sqrt(1−ρ²), and w_m=−[m(1+cos u)+1/2]/(2w). Splitting the interval at m₀ and using the exact square form for ρ yields h_m<0. The implementation checks the algebra symbolically and verifies the rational inequality chain. Consequently each parameter has one normalised root and the implicit function theorem supplies a real-analytic branch.")
    doc.add_heading("S2.5. Route 1 certificates", level=2)
    doc.add_paragraph(
        "The independent Route 1 representation uses a square 18-variable system with a gauge row. A simplified-Newton map is certified on 1,510 rational parameter boxes. Each record contains rational interval endpoints, an 18-vector rational centre, an 18×18 rational approximate inverse, a canonical SHA-256 payload, and informational decimal summaries. The verifier recomputes alpha, beta, operator bounds, contraction radius, and gap-free coverage without trusting the decimal fields.")

    doc.add_heading("S3. Goldberg-snark derivation", level=1)
    doc.add_heading("S3.1. Orbit correction", level=2)
    doc.add_paragraph(
        "The cyclic action has twelve, not fifteen, edge orbits. The apparent singleton seam edges in an exploratory script arose because canonicalisation was performed independently for differently oriented representatives. Direct orbit generation by repeated block shift gives twelve orbits of size k.")
    doc.add_heading("S3.2. Unit-chord obstruction", level=2)
    doc.add_paragraph(
        "For a rotation of angle α about the z-axis, the maximum chord ||(R−I)v|| over unit v is 2|sin(α/2)|. The v_5 channel forces a chord of length one. Thus α=2π/k is impossible for k≥7. The index (k−1)/2 gives α=π−π/k and a sufficiently long chord for every odd k≥5.")
    doc.add_heading("S3.3. Local-channel lemma", level=2)
    doc.add_paragraph("For W_y=(−yx,z_y) in the xz-plane and the quarter-turn J, let")
    add_equation(doc, "U_y=[W_y−δ_y JW_y]/[2(1−y²)].")
    doc.add_paragraph(
        "Using ||W_y||²=1−y² and δ_y²=4(1−y²)−1 gives ||U_y||=1 and U_y·W_y=1/2. Expanding U_y yields (A_y,D_y), from which the local vector identities and unit norms follow by direct calculation.")
    doc.add_heading("S3.4. Exact interval proof", level=2)
    doc.add_paragraph(
        "The certificate partitions x∈[0,13/40] into 16 boxes for each endpoint and the (s,x)-rectangle into 32×32 boxes for the derivative. Square roots are enclosed by dyadic rationals using integer square roots at 112 bits. The independent verifier checks all endpoint upper/lower signs, every derivative lower bound, positivity of all radicands, schema version, and the canonical SHA-256 digest.")
    add_figure(doc, FIG/'goldberg_scalar_branch.png', "Figure S1. Floating-point visualisation of the unique Goldberg scalar branch; exact existence and uniqueness are established by the rational certificate.", 5.5)

    doc.add_heading("S4. Extended computational validation", level=1)
    table=doc.add_table(rows=1,cols=3); table.style='Table Grid'
    for i,t in enumerate(['Campaign','Observed output','Interpretation']): table.rows[0].cells[i].text=t; set_cell_shading(table.rows[0].cells[i],'D9EAF7')
    data=[
        ('Goldberg parameter sweep','499 odd k from 5 through 1001','finite-precision formula regression'),
        ('Goldberg full expansions','149 graphs through k=301','seam, orientation, and incidence regression'),
        ('Goldberg G_1001','8,008 vertices; 12,012 edges','large-instance scalability check'),
        ('Flower exact dyadic set','J_5 through J_41','independent exact finite certificates'),
        ('Flower equivariance diagnostic','all generators for J_5,…,J_13','stored unconstrained solutions are not near the symmetric branch'),
        ('General cubic campaign','named families and random cubic graphs','finite evidence only for the universal conjecture'),
    ]
    for row in data:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    doc.add_paragraph(
        "The worst residuals from the recorded Goldberg sweep are approximately 3.04×10⁻¹⁶ for reduced Kirchhoff conservation, 2.22×10⁻¹⁵ for reduced norms, 7.20×10⁻¹⁴ for full-graph Kirchhoff conservation, and 3.71×10⁻¹⁴ for full-graph norms. These values are consistent with accumulated floating-point error and do not enter the exact proof.")

    doc.add_heading("S5. Code-to-claim map", level=1)
    map_rows=[
        ('Flower analytic theorem','scripts/stage2_reduced.py; scripts/verify_route2_algebra.py'),
        ('Flower strict-ansatz obstruction','scripts/strict_ansatz_exact.py'),
        ('Flower exact dyadic certificates','scripts/exact_rational_certificates.py; scripts/verify_exact_certificates.py'),
        ('Flower Route 1 boxes','scripts/route1_kantorovich_family.py; scripts/recheck_route1_certificates.py'),
        ('Goldberg explicit formulas','src/goldberg_s2/algebra.py; construction.py'),
        ('Goldberg exact interval theorem','src/goldberg_s2/interval*.py; scripts/run_interval_proof.py; verify_interval_certificate.py'),
        ('Full graph verification','src/goldberg_s2/graph.py; verify.py; scripts/run_numerical_sweep.py'),
        ('Complete command catalogue','RUNBOOK.md'),
    ]
    table=doc.add_table(rows=1,cols=2);table.style='Table Grid'
    for i,t in enumerate(['Claim or artefact','Primary implementation']):table.rows[0].cells[i].text=t;set_cell_shading(table.rows[0].cells[i],'D9EAF7')
    for row in map_rows:
        cells=table.add_row().cells;cells[0].text=row[0];cells[1].text=row[1]

    doc.add_heading("S6. Repository audit and recommendations", level=1)
    doc.add_paragraph(
        "The submitted archive was mathematically usable but not release-clean. The reviewed package resolves or documents the following issues.")
    issues=[
        "The original LICENSE contained an unresolved merge conflict between MIT and AGPL-3.0. The active release uses MIT, consistent with the supplied citation metadata; the conflicting original is preserved under archive/.",
        "The README referred to absent requirements files. Compatibility files have been added while pyproject extras remain the authoritative dependency declaration.",
        "Some scripts used /home/claude paths and one equivariance driver pointed to a non-existent certificate directory. Defaults are now repository-relative.",
        "Generated Python caches and .egg-info metadata were removed from the source release.",
        "The repository combines a current family-theorem package with a broader historical framework. A future major version should move family-specific modules under src/s2flow/families while retaining compatibility imports.",
        "No Git remote was present in the archive and the public repository could not be identified unambiguously. The exact GitHub URL must be added to CITATION.cff before release.",
        "A tagged GitHub release and Zenodo DOI are recommended once the manuscript and certificates are frozen.",
        "Independent mathematical review remains essential. Reproducible code reduces verification cost but does not replace referee scrutiny."
    ]
    for issue in issues:
        p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(0.2);p.paragraph_format.first_line_indent=Inches(-0.12);p.add_run('• ');p.add_run(issue)

    doc.add_heading("S7. Complete execution guide", level=1)
    doc.add_paragraph(
        "Every supported command, input, output path, and proof status is documented in RUNBOOK.md. That Markdown file is part of the supplementary reproducibility record and should be deposited with the code release.")

    doc.add_heading("S8. Limitations and open directions", level=1)
    doc.add_paragraph(
        "The two family proofs do not establish the universal S²-flow conjecture. Their reusable contribution is a methodology for periodic cubic graphs: compute free edge orbits; test representations of the automorphism group; eliminate impossible representations by local geometric constraints; solve local channels symbolically; and reduce closure to a certifiable low-dimensional system. Applying this programme to non-cyclic or multi-parameter snark families may require transfer matrices, representation decompositions with several irreducible components, or interval continuation in more than one parameter.")

    refs=[
        "R. Isaacs, Infinite families of nontrivial trivalent graphs which are not Tait colorable, American Mathematical Monthly 82 (1975), 221–239.",
        "M. K. Goldberg, Construction of class 2 graphs with maximum vertex degree 3, Journal of Combinatorial Theory, Series B 31 (1981), 282–291.",
        "H. Fleischner, B. Bagheri Gh., and B. Klocker, Perfect pseudo-matchings in cubic graphs, Graphs and Combinatorics 40 (2024), Article 118.",
        "C. Thomassen, Group flow, complex flow, unit vector flow, and the (2+epsilon)-flow conjecture, Journal of Combinatorial Theory, Series B 108 (2014), 81–91.",
        "D. Mattiolo, G. Mazzuoccolo, J. Rajník, and G. Tabarelli, On d-dimensional nowhere-zero r-flows on a graph, European Journal of Mathematics 9 (2023), Article 101.",
        "H. Houdrouge, B. Miraftab, and P. Morin, 2-dimensional unit vector flows, arXiv:2602.21526 (2026).",
    ]
    add_references(doc,refs)
    path=OUT/'S2_Flows_Flower_Goldberg_Snarks_Supplementary.docx'
    doc.save(path)
    return path


if __name__ == '__main__':
    OUT.mkdir(parents=True,exist_ok=True)
    print(build_main())
    print(build_supplement())
