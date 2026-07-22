from __future__ import annotations

from pathlib import Path
import json
import math

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'S2_flow_partial_theorems_and_structural_lemmas_ru.docx'
ASSETS = ROOT / 'docs' / 'assets_partial_results'
ASSETS.mkdir(parents=True, exist_ok=True)
RESULTS = ROOT / 'results' / 'massive_run'


def make_figures() -> tuple[Path, Path, Path]:
    # Figure 1: exact cut factorisation schematic.
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.axis('off')
    ax.set_xlim(0.0, 1.0)
    ax.set_ylim(0.0, 1.0)
    ax.text(0.25, 0.93, '2-rib incision', ha='center', va='center', fontsize=13, fontweight='bold')
    ax.text(0.75, 0.93, '3-rib incision', ha='center', va='center', fontsize=13, fontweight='bold')

    # Two-cut panel.
    ax.add_patch(plt.Circle((0.12, 0.55), 0.09, fill=False, linewidth=1.5))
    ax.add_patch(plt.Circle((0.38, 0.55), 0.09, fill=False, linewidth=1.5))
    ax.plot([0.20, 0.30], [0.60, 0.60], linewidth=1.5)
    ax.plot([0.20, 0.30], [0.50, 0.50], linewidth=1.5)
    ax.text(0.12, 0.55, 'A', ha='center', va='center', fontsize=12)
    ax.text(0.38, 0.55, 'B', ha='center', va='center', fontsize=12)
    ax.text(0.25, 0.37, 'boundary vectors: p and -p', ha='center', fontsize=10)
    ax.annotate('', xy=(0.19, 0.20), xytext=(0.31, 0.20), arrowprops={'arrowstyle': '<->'})
    ax.text(0.25, 0.12, 'closing each side with one edge', ha='center', fontsize=9)

    # Three-cut panel.
    ax.add_patch(plt.Circle((0.62, 0.55), 0.09, fill=False, linewidth=1.5))
    ax.add_patch(plt.Circle((0.88, 0.55), 0.09, fill=False, linewidth=1.5))
    for y in (0.63, 0.55, 0.47):
        ax.plot([0.70, 0.80], [y, y], linewidth=1.5)
    ax.text(0.62, 0.55, 'A', ha='center', va='center', fontsize=12)
    ax.text(0.88, 0.55, 'B', ha='center', va='center', fontsize=12)
    ax.text(0.75, 0.34, 'p₁+p₂+p₃=0,  pᵢ·pⱼ=-1/2', ha='center', fontsize=10)
    ax.annotate('', xy=(0.69, 0.20), xytext=(0.81, 0.20), arrowprops={'arrowstyle': '<->'})
    ax.text(0.75, 0.12, 'closing each side with a new cubic vertex', ha='center', fontsize=9)
    fig.tight_layout()
    cut_path = ASSETS / 'cut_factorization.png'
    fig.savefig(cut_path, dpi=220, bbox_inches='tight')
    plt.close(fig)

    df = pd.read_csv(RESULTS / 'summary.csv')
    counts = df['method'].value_counts()
    fig, ax = plt.subplots(figsize=(7.2, 4.5))
    labels = [''Exact Construction\nfrom 3-Coloring', 'Nonlinear Search\nin the Space of Cycles'']
    values = [int(counts.get('exact_three_edge_colouring_construction', 0)), int(counts.get('cycle_space_nonlinear_least_squares', 0))]
    bars = ax.bar(labels, values)
    ax.set_ylabel('Number of graphs')
    ax.set_title('Types of certificates in a reproducible campaign')
    ax.grid(axis='y', alpha=0.25)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, value + max(values)*0.015, str(value), ha='center', va='bottom', fontsize=11)
    fig.tight_layout()
    cert_path = ASSETS / 'certificate_counts.png'
    fig.savefig(cert_path, dpi=220, bbox_inches='tight')
    plt.close(fig)

    nonlinear = df[df['method'] == 'cycle_space_nonlinear_least_squares'].copy()
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.scatter(range(len(nonlinear)), nonlinear['max_unit_norm_residual'])
    ax.set_yscale('log')
    ax.set_xticks(range(len(nonlinear)))
    ax.set_xticklabels(nonlinear['name'], rotation=35, ha='right', fontsize=8)
    ax.set_ylabel('Maximum error of the norm')
    ax.set_title('Independently verified numerical certificates')
    ax.grid(axis='y', alpha=0.25)
    fig.tight_layout()
    residual_path = ASSETS / 'nonlinear_residuals.png'
    fig.savefig(residual_path, dpi=220, bbox_inches='tight')
    plt.close(fig)
    return cut_path, cert_path, residual_path


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        for r in (r1, r2):
            r.font.name = 'Arial'
            r.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
    return p


def add_theorem(doc: Document, title: str, statement: str, proof_paragraphs: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    statement_p = add_body(doc, statement)
    statement_p.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    r = p.add_run('Proof.')
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    if proof_paragraphs:
        r2 = p.add_run(' ' + proof_paragraphs[0])
        r2.font.name = 'Arial'
        r2.font.size = Pt(10.5)
    for paragraph in proof_paragraphs[1:]:
        add_body(doc, paragraph)
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = q.add_run('□')
    rr.font.name = 'Cambria Math'
    rr.font.size = Pt(11)


def build_doc() -> None:
    cut_fig, cert_fig, residual_fig = make_figures()
    campaign = json.loads((RESULTS / 'campaign.json').read_text(encoding='utf-8'))
    df = pd.read_csv(RESULTS / 'summary.csv')
    nonlinear = df[df['method'] == 'cycle_space_nonlinear_least_squares'].copy()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(1.8)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10.5)
    for style_name, size in [('Title', 20), ('Heading 1', 15), ('Heading 2', 12), ('Heading 3', 11)]:
        styles[style_name].font.name = 'Arial'
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor(31, 55, 77)

    header = sec.header.paragraphs[0]
    header.text = 'S²-flows on bridgeless cubic graphs'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 90, 90)
    add_page_number(sec.footer.paragraphs[0])

    title = doc.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('New partial theorems and structural lemmas\nfor S²-flows')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run('Complete proofs, boundaries of novelty, and reproducible computational tests')
    r.italic = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = date_p.add_run('State of the art and computation: July 22, 2026')
    rr.font.name = 'Arial'
    rr.font.size = Pt(9)
    doc.add_paragraph()

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    shade_cell(cell, 'EAF2F8')
    set_cell_text(cell, 'The main result of the paper: the complete S²-flow conjecture is not proven. Factorization theorems, constructive special cases, and reductions of the minimal counterexample, new to this project, are obtained and rigorously proven. The claim of world scientific novelty requires peer review and additional bibliographic research..', bold=True, size=10)

    doc.add_heading('1. Summary of results', level=1)
    summary_items = [
        'Exact cut law: the sum of signed vectors on any edge cut is zero.'
        'Small cut rigidity: on a 2-cut, vectors are antipodal; on a 3-cut, they form an equilateral triple with pairwise scalar products of -1/2.'
        'Exact factorization by a 2-edge cut: a graph has an S²-flow if and only if both canonically closed sides have S²-flows.'
        'Exact factorization by a nontrivial 3-edge cut with the same "if and only if" condition.'
        'The injection of one cubic graph into a vertex of another admits the reverse reduction. Inflating a vertex into a triangle preserves the S²-flow in both directions.'
        'Every 3-edge-colorable cubic graph has an explicit S²-flow constructed from three directed 2-factors.'
        'Every planar bridgeless cubic graph has an S²-flow as a consequence of the four-color theorem and the construction from a 3-edge-coloring.'
        'A minimal counterexample, if it exists, must be cyclically 4-edge-connected and triangle-free.'
        'A reproducible code was constructed that yielded 288 of 288 independently verified finite certificates.'
    ]
    for item in summary_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_heading('2. Initial Statement and Status of the Open Problem', level=1)
    add_body(doc, 'Let G=(V,E) be a finite directed graph. An S²-flow is a mapping φ:E→R³ such that ||φ(e)||₂=1 on each edge and Kirchhoff's law holds at each vertex. Changing the edge orientation is accompanied by replacing φ(e) with -φ(e).')
    add_equation(doc, 'Σₑ ε(v,e) φ(e) = 0 for every v∈V, ||φ(e)||₂ = 1 for every e∈E.')
    add_body(doc, 'Jains conjecture states that every bridgeless graph, and equivalently every bridgeless cubic graph, admits such a flow. As of July As of 2026, the conjecture remains open. Houdrouge, Miraftab, and Morin (2026) provided a geometric equivalence via equiangular spherical immersions, a one-sided composition theorem, preservation under triangle inflation, and explicit quasi-Petersen families. Mattiolo et al. (2023) captures the general higher-dimensional context and a known universal bound in higher dimensions.')

    doc.add_heading('3. Статус новизны', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Результат', 'Статус', 'Комментарий']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], 'D9EAF7')
    novelty_rows = [
        ('Геометрическая характеризация S²-потока', 'Известно', 'Houdrouge, Miraftab, Morin, 2026.'),
        ('Композиция двух графов с потоками', 'Известно', 'Односторонняя теорема 17 в работе 2026 года.'),
        ('Раздувание вершины в треугольник', 'Известно в прямую сторону', 'Теорема 19 в работе 2026 года.'),
        ('Точная 2-cut факторизация', 'Получено в проекте', 'Полное доказательство приведено ниже; точный приоритет не гарантируется.'),
        ('Точная 3-cut факторизация и обратная композиция', 'Получено в проекте', 'Усиливает одностороннюю композицию до эквивалентности.'),
        ('Инвариантность при triangle blow-up', 'Получено в проекте', 'Обратное направление следует из 3-cut факторизации.'),
        ('Явный поток из 3-edge-colouring', 'Полный частный результат', 'Следует также из общего принципа цикл-покрытий; здесь дана автономная конструкция.'),
        ('Редукция минимального контрпримера', 'Получено как следствие', 'Циклическая 4-связность и отсутствие треугольников.')
    ]
    for row in novelty_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)

    doc.add_heading('4. Базовые структурные леммы', level=1)
    add_theorem(doc, 'Лемма 1. Закон рёберного разреза.', 'Для любого множества вершин U⊂V и любого S²-потока сумма векторов на рёбрах δ(U), подписанных наружу из U, равна нулю.', [
        'Просуммируем законы Кирхгофа по всем вершинам U. Каждое внутреннее ребро имеет два конца в U и входит в сумму дважды с противоположными знаками, поэтому сокращается. Остаются только рёбра δ(U), каждое ровно один раз с наружным знаком. Их сумма равна нулю.'
    ])
    add_theorem(doc, 'Лемма 2. Жёсткость единичной тройки.', 'Если u,v,w∈R³, ||u||=||v||=||w||=1 и u+v+w=0, то u·v=v·w=w·u=-1/2. Следовательно, три вектора лежат в одной плоскости и образуют углы 2π/3.', [
        'Из w=-(u+v) и ||w||²=1 получаем 1=||u+v||²=2+2u·v, откуда u·v=-1/2. Циклическая перестановка даёт два остальных равенства. Линейная зависимость u+v+w=0 обеспечивает компланарность.'
    ])
    add_body(doc, 'Следствие. На 2-рёберном разрезе два наружных единичных вектора равны p и -p. На 3-рёберном разрезе наружные векторы образуют единственную с точностью до ортогонального преобразования равностороннюю тройку.')
    doc.add_picture(str(cut_fig), width=Inches(6.7))
    cap = doc.add_paragraph('Рисунок 1. Канонические интерфейсы 2- и 3-рёберных разрезов.')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

    doc.add_heading('5. Точная факторизация по малым разрезам', level=1)
    add_theorem(doc, 'Теорема 3. Точная факторизация по 2-рёберному разрезу.', 'Пусть G - бесмостовый кубический граф и δ(A)={a₁b₁,a₂b₂} - нетривиальный 2-рёберный разрез, где aᵢ∈A и bᵢ∉A. Построим G_A, удалив разрез и добавив ребро a₁a₂; аналогично построим G_B. Тогда G имеет S²-поток тогда и только тогда, когда G_A и G_B имеют S²-потоки.', [
        'Необходимость. Пусть φ - поток на G. Подпишем два граничных вектора наружу из A как p₁ и p₂. По Лемме 1 p₁+p₂=0. После удаления разреза добавим ребро a₁a₂ и ориентируем его так, чтобы его вклады в a₁ и a₂ были соответственно p₁ и p₂. Это возможно, поскольку p₂=-p₁ и ||p₁||=1. Во всех остальных вершинах уравнения не меняются. Получен поток на G_A. То же построение применяется к B.',
        'Достаточность. Пусть на G_A и G_B заданы потоки. Замыкающие рёбра несут единичные векторы x и y. Применим ко всему потоку на G_B ортогональное преобразование Q, переводящее y в x или -x в соответствии с выбранным попарным соединением концов. Удалим замыкающие рёбра и восстановим два рёбра разреза, присвоив им x с ориентациями, воспроизводящими прежние вклады в четыре граничные вершины. На остальных вершинах баланс сохранён, а на граничных вклады совпадают с удалёнными замыкающими рёбрами.'
    ])
    add_theorem(doc, 'Теорема 4. Точная факторизация по нетривиальному 3-рёберному разрезу.', 'Пусть δ(A)={aᵢbᵢ: i=1,2,3}. Замкнём сторону A новой вершиной z_A, соединённой с a₁,a₂,a₃; аналогично построим G_B. Тогда G имеет S²-поток тогда и только тогда, когда G_A и G_B имеют S²-потоки.', [
        'Необходимость. Наружные из A граничные векторы p₁,p₂,p₃ удовлетворяют p₁+p₂+p₃=0. Добавим новую вершину z_A и три ребра z_Aaᵢ так, чтобы вклад нового ребра в aᵢ был pᵢ. В z_A сумма вкладов равна -(p₁+p₂+p₃)=0. Нормы равны единице. Аналогично замыкается сторона B.',
        'Достаточность. В потоках на G_A и G_B тройки векторов при новых вершинах являются равносторонними по Лемме 2. Для любой заданной биекции между тремя граничными рёбрами существует ортогональное преобразование Q∈O(3), переводящее упорядоченную тройку одной стороны в противоположную тройку другой. Применим Q ко всему потоку на G_B, удалим z_A,z_B и соединим соответствующие граничные вершины. Новые рёбра воспроизводят удалённые вклады, поэтому закон Кирхгофа сохраняется.'
    ])
    add_body(doc, 'Лемма о бесмостовости замыканий. Если исходный G бесмостов, то канонические замыкания нетривиальных 2- и 3-разрезов также бесмостовы. Действительно, внутреннее ребро, являющееся мостом замыкания, отделяло бы компоненту без граничных терминалов и было бы мостом исходного графа; новое замыкающее ребро или ребро к новой вершине не может быть мостом, поскольку соответствующий терминал имеет внутренний путь к остальным терминалам, иначе исходное разрезное ребро было бы мостом.')

    doc.add_heading('6. Усиления композиционных результатов', level=1)
    add_theorem(doc, 'Следствие 5. Обратимость инъекции.', 'Пусть H▷G получен инъекцией кубического графа H в вершину кубического графа G. Тогда H▷G имеет S²-поток тогда и только тогда, когда оба графа H и G имеют S²-потоки.', [
        'Прямое направление совпадает с известной композицией: равносторонние тройки в удаляемых вершинах совмещаются ортогональным преобразованием. Обратное направление следует из Теоремы 4: три соединяющих ребра образуют нетривиальный 3-разрез, а его канонические замыкания изоморфны H и G.'
    ])
    add_theorem(doc, 'Следствие 6. Инвариантность раздувания вершины в треугольник.', 'Пусть G△ получен из кубического графа G заменой вершины треугольником, каждая вершина которого наследует одно внешнее ребро. Тогда G△ имеет S²-поток тогда и только тогда, когда G имеет S²-поток.', [
        'Граф G△ является инъекцией K₄ в соответствующую вершину G. Граф K₄ имеет S²-поток. Прямое направление следует из композиции, обратное - из Следствия 5. Таким образом, известное сохранение при раздувании усиливается до эквивалентности.'
    ])

    doc.add_heading('7. Полный конструктивный частный случай', level=1)
    add_theorem(doc, 'Теорема 7. Все 3-рёберно раскрашиваемые кубические графы.', 'Каждый кубический граф с правильной 3-рёберной раскраской допускает S²-поток, причём поток строится явно за полиномиальное время после получения раскраски.', [
        'Пусть M₁,M₂,M₃ - три цветовых класса. Для i=1,2,3 положим Hᵢ=E(G)\\Mᵢ. В каждой вершине ровно два ребра Hᵢ, поэтому Hᵢ является 2-фактором, то есть дизъюнктным объединением циклов. Ориентируем каждый цикл каждого Hᵢ произвольно.',
        'Зафиксируем глобальную ориентацию рёбер G. Для ребра e определим σᵢ(e)=0, если e∉Hᵢ; σᵢ(e)=1, если ориентация e в Hᵢ совпадает с глобальной; и σᵢ(e)=-1 в противном случае. Каждый столбец σᵢ является скалярной циркуляцией, поскольку он является суммой ориентированных циклов.',
        'Определим φ(e)=(σ₁(e),σ₂(e),σ₃(e))/√2. Координатно закон Кирхгофа выполняется. Каждое ребро принадлежит ровно двум из H₁,H₂,H₃, поэтому в φ(e) ровно две координаты равны ±1/√2, а третья равна нулю. Следовательно, ||φ(e)||²=(1+1)/2=1.'
    ])
    add_body(doc, 'Следствие 7.1. Каждый планарный бесмостовый кубический граф имеет S²-поток. По теореме Тэйта, эквивалентной теореме о четырёх красках, такой граф имеет правильную 3-рёберную раскраску; затем применяется Теорема 7.')
    add_body(doc, 'Следствие 7.2. Каждый кубический граф нечётности 0 имеет S²-поток. Чётный 2-фактор раскрашивается попеременно двумя цветами, а дополнительное совершенное паросочетание получает третий цвет.')

    doc.add_heading('8. Новый рекурсивный бесконечный класс', level=1)
    add_theorem(doc, 'Теорема 8. Класс атомов и малых сумм.', 'Пусть 𝒞 - наименьший класс кубических графов, содержащий все 3-рёберно раскрашиваемые кубические графы и известное семейство quasi-Petersen Gₐ,ᵦ,ₚ при p/6<a,b<p/2, и замкнутый относительно 2-рёберной суммы и инъекции по вершине. Тогда каждый граф из 𝒞 имеет S²-поток.', [
        'Базовые 3-рёберно раскрашиваемые графы покрываются Теоремой 7. Семейство quasi-Petersen имеет равноугловую S²-иммерсию по Proposition 2 работы Houdrouge, Miraftab и Morin, следовательно имеет S²-поток.',
        'Индукционный шаг по 2-рёберной сумме следует из достаточности Теоремы 3. Индукционный шаг по инъекции следует из достаточности Теоремы 4. Поэтому любое конечное дерево таких операций сохраняет существование S²-потока.'
    ])
    add_body(doc, 'Класс 𝒞 содержит, в частности, все итерационные раздувания вершин в треугольники, все композиции Petersen/quasi-Petersen блоков с 3-рёберно раскрашиваемыми блоками, а также графы с произвольным деревом нетривиальных 2- и 3-разрезов, если каждый атом дерева принадлежит базовому классу.')

    doc.add_heading('9. Структура минимального контрпримера', level=1)
    add_theorem(doc, 'Теорема 9. Редукция минимального контрпримера.', 'Если S²-flow conjecture ложна и G является контрпримером с минимальным числом вершин среди бесмостовых кубических графов, то G не имеет 2-рёберных разрезов, не имеет нетривиальных 3-рёберных разрезов и не содержит треугольников. В частности, G циклически 4-рёберно связен и имеет обхват не менее 4.', [
        'Если существует нетривиальный 2-разрез, обе канонические стороны меньше G и бесмостовы. По минимальности они имеют S²-потоки, а Теорема 3 склеивает их в поток на G, противоречие.',
        'Если существует нетривиальный 3-разрез, обе стороны после добавления новой вершины меньше G и бесмостовы. По минимальности они имеют потоки, а Теорема 4 снова даёт противоречие. Следовательно, разрез менее четырёх рёбер не может разделять две циклические части.',
        'Пусть G содержит треугольник. После уже доказанного отсутствия 2-разрезов три внешних соседа треугольника различны. Сожмём треугольник в одну кубическую вершину и получим меньший бесмостовый кубический граф G₀. По минимальности G₀ имеет S²-поток. По инвариантности triangle blow-up исходный G также имеет поток, противоречие.'
    ])
    add_body(doc, 'Замечание. Полученная редукция не доказывает гипотезу, но исключает все «составные» препятствия и переносит поиск на циклически 4-рёберно связные треугольник-свободные атомы. Для строгого определения snark с обхватом не менее 5 остаётся отдельно устранить 4-циклы; универсальная 4-cycle reduction в данной работе не доказана.')

    doc.add_heading('10. Алгебраическая структурная лемма', level=1)
    add_theorem(doc, 'Лемма 10. Rank-3 Gram compression.', 'Пусть B - ориентированная матрица инцидентности, Z - матрица, столбцы которой образуют базис ker B, и zₑ - строка Z для ребра e. Если существует Q⪰0 с rank Q≤3 и zₑQzₑᵀ=1 для каждого e, то G имеет S²-поток.', [
        'Разложим Q=AAᵀ, где A имеет не более трёх столбцов; при необходимости дополним нулевыми столбцами до трёх. Положим X=ZA. Тогда BX=BZA=0, поскольку BZ=0. Для каждой строки xₑ имеем ||xₑ||²=zₑAAᵀzₑᵀ=zₑQzₑᵀ=1. Следовательно, строки X задают S²-поток.'
    ])
    add_body(doc, 'Эта лемма разделяет задачу на выпуклую часть Q⪰0 с линейными диагональными ограничениями и невыпуклое условие rank Q≤3. Код проверяет такие сертификаты независимо от нелинейного оптимизатора.')

    doc.add_heading('11. Масштабная воспроизводимая проверка', level=1)
    add_body(doc, f"Проведена детерминированная кампания на {campaign['instances']} графах: 270 случайных бесмостовых кубических графов порядков 10, 12, 14, 16, 18, 20, 24, 30 и 40, по 30 графов каждого порядка, а также 18 именованных и параметрических графов. Seed: {campaign['config']['seed']}.")
    add_body(doc, f"Получено {campaign['exact_constructive_certificates']} точных конструктивных сертификатов из 3-рёберных раскрасок и {campaign['nonlinear_certificates']} численных rank-3 сертификатов для нераскрашиваемых или специально выбранных графов. Независимая повторная проверка прочитала все NPZ-сертификаты заново и подтвердила {campaign['valid_certificates']} из {campaign['instances']}.")
    add_body(doc, f"Максимальная ошибка закона Кирхгофа: {campaign['maximum_conservation_residual']:.3e}. Максимальная ошибка единичной нормы среди численных решений: {campaign['maximum_unit_norm_residual']:.3e}. Для численных сертификатов применён независимый порог 10⁻⁷; конструктивные сертификаты имеют ошибки порядка машинного округления.")
    doc.add_picture(str(cert_fig), width=Inches(5.8))
    p = doc.add_paragraph('Рисунок 2. Распределение типов сертификатов.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    doc.add_picture(str(residual_fig), width=Inches(6.5))
    p = doc.add_paragraph('Рисунок 3. Ошибки нормы для восьми нелинейных сертификатов.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Граф', '|V|', '3-раскраска', 'Метод', 'Ошибка баланса', 'Ошибка нормы']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
        shade_cell(table.rows[0].cells[i], 'D9EAF7')
    for _, row in nonlinear.iterrows():
        cells = table.add_row().cells
        values = [
            str(row['name']),
            str(int(row['vertices'])),
            str(row['three_edge_coloring_status']),
            'cycle-space LSQ',
            f"{row['max_conservation_residual']:.2e}",
            f"{row['max_unit_norm_residual']:.2e}",
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, size=7.5)

    doc.add_heading('12. Reproducibility', level=1)
    add_body(doc, 'Основной запуск:')
    code = doc.add_paragraph()
    code.style = doc.styles['Normal']
    run = code.add_run('python scripts/massive_verification.py --orders 10 12 14 16 18 20 24 30 40 --samples-per-order 30 --nonlinear-restarts 16 --output-dir results/massive_run')
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    add_body(doc, 'Независимая проверка сохранённых сертификатов:')
    code = doc.add_paragraph()
    run = code.add_run('python scripts/verify_campaign.py results/massive_run --tolerance 1e-7')
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    add_body(doc, 'Для исчерпывающих потоков graph6 из nauty предусмотрен streaming-скрипт:')
    code = doc.add_paragraph()
    run = code.add_run('geng -c -d3 -D3 16 | python scripts/nauty_cubic_campaign.py --output-dir results/geng16')
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    add_body(doc, 'Каждый сертификат содержит graph6-представление графа, стабильный порядок рёбер и матрицу X. Проверяющий скрипт заново строит матрицу инцидентности и вычисляет BX, нормы строк и ранг Gram-матрицы. Успех конечной проверки не является доказательством бесконечного утверждения.')

    doc.add_heading('13. Ограничения и следующий доказательный барьер', level=1)
    limitations = [
        'Полная гипотеза не доказана.',
        'Точные теоремы по разрезам сводят проблему к циклически 4-рёберно связным атомам, но не решают эти атомы.',
        'Численные решения для снарков являются проверяемыми приближёнными сертификатами, а не символьными доказательствами существования для бесконечного семейства.',
        'Не получена универсальная редукция 4-цикла, которая позволила бы автоматически повысить обхват минимального контрпримера до 5.',
        'Точный мировой приоритет факторизационных формулировок должен быть подтверждён независимым библиографическим обзором и рецензированием.'
    ]
    for item in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_heading('14. Вывод', level=1)
    add_body(doc, 'Получены полные доказательства нескольких содержательных частных результатов. Наиболее сильное структурное продвижение состоит в точной двусторонней факторизации S²-потоков по 2- и 3-рёберным разрезам. Она превращает композиционные конструкции в теорему разложения, усиливает triangle blow-up до эквивалентности и показывает, что минимальный контрпример должен быть атомарным. Конструктивная теорема для всех 3-рёберно раскрашиваемых кубических графов покрывает планарный случай и даёт точные машинно проверяемые сертификаты. Открытым ядром остаются циклически 4-рёберно связные нераскрашиваемые кубические графы.')

    doc.add_heading('Литература', level=1)
    refs = [
        'H. Houdrouge, B. Miraftab, P. Morin. 2-Dimensional Unit Vector Flows. arXiv:2602.21526, 2026.',
        'D. Mattiolo, G. Mazzuoccolo, J. Rajník, G. Tabarelli. On d-dimensional nowhere-zero r-flows on a graph. European Journal of Mathematics 9, Article 101, 2023. DOI: 10.1007/s40879-023-00694-1.',
        'C. Thomassen. Group flow, complex flow, unit vector flow, and the (2+epsilon)-flow conjecture. Journal of Combinatorial Theory, Series B 108, 81-91, 2014. DOI: 10.1016/j.jctb.2014.02.012.',
        'G. Brinkmann, J. Goedgebeur, J. Hägglund, K. Markström. Generation and properties of snarks. Journal of Combinatorial Theory, Series B 103(4), 468-488, 2013.',
        'R. Diestel. Graph Theory, 5th edition. Springer, 2018.',
        'W. T. Tutte. A contribution to the theory of chromatic polynomials. Canadian Journal of Mathematics 6, 80-91, 1954.'
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Number')
        p.add_run(ref)

    doc.core_properties.title = 'Новые частные теоремы и структурные леммы для S²-потоков'
    doc.core_properties.subject = 'S2-flow conjecture, cubic graphs, structural decomposition, reproducibility'
    doc.core_properties.author = 'Research synthesis'
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build_doc()
